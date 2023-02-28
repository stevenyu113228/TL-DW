import cv2
import os
import pytesseract
import openai
import sys
import speech_recognition as sr
from sewar.full_ref import uqi
from pathlib import Path
from multiprocessing import Pool
from pydub import AudioSegment
from tqdm import tqdm
from PIL import Image
from docx import Document
from docx.shared import Inches
from datetime import timedelta


OPEN_AI_KEY = "" 
THRESHOLD = 0.90 # 相似度門檻
THREADS = 8 # 8個執行緒

def start(file_name):
    # 建立資料夾、轉檔
    Path(f'tmp/{file_name.stem}').mkdir(parents=True, exist_ok=True)
    Path(f'tmp/{file_name.stem}/img').mkdir(parents=True, exist_ok=True)
    Path(f'tmp/{file_name.stem}/audio').mkdir(parents=True, exist_ok=True)
    Path(f'tmp/{file_name.stem}/a2t').mkdir(parents=True, exist_ok=True)
    Path(f'tmp/{file_name.stem}/i2t').mkdir(parents=True, exist_ok=True)
    Path(f'tmp/{file_name.stem}/ai').mkdir(parents=True, exist_ok=True)

    os.system(f'ffmpeg -i {file_name} -vf fps=1 tmp/{file_name.stem}/img/%d.png')
    os.system(f'ffmpeg -i {file_name} -ac 2 -f wav tmp/{file_name.stem}/out.wav')

def get_files():
    # 取得所有檔案 & 排序
    files = os.listdir(f'tmp/{file_name.stem}/img')
    files.sort(key=lambda f: int(''.join(filter(str.isdigit, f))))
    files = [os.path.join(f'tmp/{file_name.stem}/img', f) for f in files]
    return files

def split_files(file_names):
    # 分割檔案
    step = len(file_names) // THREADS
    groups_name = [file_names[i:i + step] for i in range(0, len(file_names), step)]
    return groups_name

def comp(files):
    # 相似度比對
    for i,n in enumerate(files):
        try:
            img1 = cv2.imread(n)
            img2 = cv2.imread(files[i+1])
            if uqi(img1, img2) > THRESHOLD:
                os.remove(n)
            else:
                print("Image", n, "is different from", files[i+1])
        except:
            pass

def insert_average(lst):
    """
    對一個有序的數列 lst，檢查每個數字與其前一個數字的差值是否超過 180，
    若是，則在兩者之間插入前後的平均值，並遞迴處理，直到所有相鄰數字之間的差值都 <= 180。
    """
    new_lst = [lst[0]]  # 用一個新的 List 儲存處理後的數列
    for i in range(1, len(lst)):
        diff = lst[i] - new_lst[-1]  # 計算相鄰兩個數字的差值
        if diff > 180:
            avg = int((lst[i] + new_lst[-1]) / 2)  # 計算前後數字的平均值
            new_lst.append(avg)  # 將平均值加入新的數列中
            new_lst.extend(insert_average([lst[i], avg]))  # 遞迴處理剩下的部分
        else:
            new_lst.append(lst[i])  # 若差值 <= 180，則直接加入新的數列中
    new_lst = list(set(new_lst))
    new_lst.sort()
    return new_lst



def split_audio(p,file_name):
    for i, j in enumerate(p):
        if i == 0:
            t1 = 0
        else:
            t1 = p[i-1] * 1000
        t2 = j * 1000
        newAudio = AudioSegment.from_wav(f"tmp/{file_name}/out.wav")
        newAudio = newAudio[t1:t2]
        newAudio.export(f'tmp/{file_name}/audio/{j}.wav', format="wav")

def audio2text(file_name):
    r = sr.Recognizer()
    for wav_file in tqdm(list(Path(f'tmp/{file_name}/audio').glob('*.wav'))):
        file = sr.AudioFile(str(wav_file))
        with file as source:
            try:
                audio = r.record(source)
                res = r.recognize_google(audio, language='en-US')
            except Exception as e:
                res = "Error"
            with open(f'tmp/{file_name}/a2t/{wav_file.stem}.txt', 'w') as f:
                f.write(res)
    

def img2text(file_name):
    for img_file in tqdm(list(Path(f'tmp/{file_name}/img').glob('*.png'))):
        img = Image.open(img_file)
        text = pytesseract.image_to_string(img, lang='eng')
        with open(f'tmp/{file_name}/i2t/{img_file.stem}.txt', 'w') as f:
            f.write(text)


def chat_ai(voice, ocr):
    try:
        openai.api_key = OPEN_AI_KEY
        response = openai.Completion.create(
            model="text-davinci-003",
            prompt=f"以下是一段課堂上投影片的語音轉文字，以及其對應講義的 OCR 資料，請忽視 OCR 的錯誤及簡報 header/footer。並使用繁體中文統整該頁面之內容，專有名詞部分請盡量保留使用英文\n\n語音轉文字：```\n{voice}\n```\n\n投影片 OCR：```{ocr}```\n\n統整結果：",
            temperature=0.7,
            max_tokens=3000,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
        )
        res = response['choices'][0]['text']
    except Exception as e:
        res = f"Error, {e}"
    return res.strip()


def ai_summary(file_name):
    transcript = list(Path(f'tmp/{file_name}/a2t').glob('*.txt'))
    transcript = [t.stem for t in transcript]
    transcript.sort(key=lambda f: int(''.join(filter(str.isdigit, f))))

    image = list(Path(f'tmp/{file_name}/i2t').glob('*.txt'))
    image = [i.stem for i in image]
    image.sort(key=lambda f: int(''.join(filter(str.isdigit, f))))


    tess = Path(f'tmp/{file_name}/i2t') / (image[0] + '.txt')
    for t in tqdm(transcript):
        tran = open(f"tmp/{file_name}/a2t/{t}.txt",'r').read()
        if os.path.exists(f"tmp/{file_name}/i2t/{t}.txt"):
            tess = open(f"tmp/{file_name}/i2t/{t}.txt",'r').read()
        ai_res = chat_ai(tran,tess)
        with open(f'tmp/{file_name}/ai/{t}.txt','w') as f:
            f.write(ai_res)

def to_docx(file_name):
    document = Document()
    document.add_heading(file_name, 0)

    transcript = list(Path(f'tmp/{file_name}/a2t').glob('*.txt'))
    transcript = [t.stem for t in transcript]
    transcript.sort(key=lambda f: int(''.join(filter(str.isdigit, f))))


    tess = transcript[0] # te
    for t in transcript:
        tran = t
        if os.path.exists(f"tmp/{file_name}/img/{t}.png"):
            tess = t
        
        document.add_picture(f'tmp/{file_name}/img/{tess}.png', width=Inches(6.0))

        lst_time = int(transcript[transcript.index(tran)-1]) if transcript.index(tran)-1 >= 0 else 0
        par = document.add_paragraph(f"{str(timedelta(seconds=lst_time))} ~ {str(timedelta(seconds=int(tran)))} ")
        par.alignment = 1 # Center

        document.add_heading('AI Summary', level=1)
        document.add_paragraph(open(f'tmp/{file_name}/ai/{tran}.txt').read())

        document.add_heading('Transcript', level=1)
        document.add_paragraph(open(f'tmp/{file_name}/a2t/{tran}.txt').read())
        document.add_page_break()

    document.save(f'{file_name}.docx')
    print("Output: ", f'{file_name}.docx')

if __name__ == '__main__':
    if OPEN_AI_KEY == "":
        print("Please set OPEN_AI_KEY in Line 18")
        exit(0)

    if len(sys.argv) == 1:
        print(f"Usage: python3 {sys.argv[0]} [file_name]")
        exit(0)

    file_name = sys.argv[1]
    file_name = Path(file_name)
    if ' ' in file_name.name:
        print("file_name can't have space")
        exit(0)

    print("[!] Start processing ...")
    start(file_name)
    file_names = get_files()
    groups_name = split_files(file_names)
    p = Pool(THREADS)
    p.map(comp, groups_name)
    
    file_names = get_files()
    comp(file_names)

    print("[!] Start split audio ...")
    p = [str(i.stem) for i in list(Path(f'tmp/{file_name.stem}/img').glob('*.png'))]
    p.sort(key=lambda f: int(''.join(filter(str.isdigit, f))))
    p = [int(i) for i in p]

    p = insert_average(p)
    split_audio(p, file_name.stem)

    print("[!] Start audio to text ...")
    audio2text(file_name.stem)

    print("[!] Start Tesseract ...")
    img2text(file_name.stem)

    print("[!] Start AI Summary ...")
    ai_summary(file_name.stem)

    # 合成 Docx
    print("[!] Start generate docx ...")
    to_docx(file_name.stem)